import docx
import PyPDF2

# This program will compare and find the differences between two contracts or docs page by page


def ReadingText(filename):  # get each paragraph
    count_pargphs = 0
    doc = docx.Document(filename)
    completed_file1 = []
    # print("paragraphs ", end=""),
    # print(len(doc.paragraphs))

    for paragraph in doc.paragraphs:
        completed_file1.append(paragraph.text)

        # if completed_file1.append(paragraph.text) is not None:
        #     count_pargphs = count_pargphs + 1
            # print(count_pargphs)
            # print(completed_file1.append(paragraph.text))

    # return '\n'.join(completed_file1)
    return completed_file1


def ReadingText2(filename2):  # get each paragraph
    doc2 = docx.Document(filename2)
    completed_file2 = []

    for paragraph in doc2.paragraphs:
        completed_file2.append(paragraph.text)

    # king = '\n'.join(completed_file2)
    return '\n'.join(completed_file2)


def CompareParagraphs(paragraph1, paragraph2):  # compare two paragraphs each paragraph
    print(paragraph1)
    print(paragraph2)
    print(" ")
    if paragraph1 == paragraph2:
        print("The paragraphs are equal")
    else:
        print("The paragraphs are not equal")


def Turn_toList(line_1):  #

    thisList = []
    for word_ in line_1.split():
        thisList.append(word_)

    return thisList

########################################################################################################################
########################################################################################################################

trial = ReadingText('file1.docx')
trial2 = ReadingText2('file2.docx')
pargrph_list = Turn_toList(trial2)
report = docx.Document()

i = 0

########################################################################################################################
##################################################  M A I N  ###########################################################
########################################################################################################################

# CompareParagraphs(line1, line2)

print("################################################ S T A R T ######################################################")
print(" ")
print("The possibly longer file needs to be File 1. The differences will be reference in regards to File 1 and they"
      "will be in bold.")
print("\nEQUALITY => File 1 (=/!=) File 2")
print(" ")

count_paragraphs = 0
for palavra in trial:
    count_ctrl = 0
    paragrafo = report.add_paragraph()
    count_ctrl += 1
    count_paragraphs += 1
    for word in palavra.split():

        for word2 in pargrph_list[i:]:

            if word == word2:
                paragrafo.add_run(word + " ")
                report.save("report.docx")
                i += 1
                break

            else:
                if count_ctrl == 1:
                    print("\033[1;34m" + "The Paragraph Number" + "\033[1;m", end=" ")
                    print(count_paragraphs, end=" ")
                    print("\033[1;34m" + "Below Is Different" + "\033[1;m")
                    count_ctrl = 0
                print(" UNEQUAL =>", end=" "),
                # print("\033[44;33m" + word + " != " + word2 + "\033[m")
                print("\033[1;32m" + word + " != " + word2 + "\033[1;m")
                paragrafo.add_run(word + " ").bold = True

                i += 1
                break

pdfFileObj = open('file1.pdf', 'rb')
# creating a pdf reader object
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

# printing number of pages in pdf file
print("Number of pager: ", end=""), print(pdfReader.numPages)